"""Resume parser for PDF and DOCX files.

The parser extracts text without OCR, identifies common resume sections, and
returns the existing ``CandidateRecord`` model so downstream matching, merging,
confidence, provenance, projection, and validation can remain unchanged.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from candidate_transformer.domain import CandidateRecord
from candidate_transformer.ingestion.base import CandidateParser
from candidate_transformer.ingestion.exceptions import ParserFileNotFoundError, ParserReadError, ParserSchemaError
from candidate_transformer.normalization import (
    DateNormalizer,
    EmailNormalizer,
    LocationNormalizer,
    PhoneNormalizer,
    SkillNormalizer,
)
from candidate_transformer.utils import LinkClassifier

logger = logging.getLogger(__name__)


class ResumeParser(CandidateParser):
    """Parse PDF and DOCX resumes into ``CandidateRecord`` instances."""

    _EMAIL_PATTERN = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.IGNORECASE)
    _PHONE_PATTERN = re.compile(r"(?:\+\d{1,3}[\s().-]*)?(?:\d[\s().-]*){7,}\d")
    _URL_PATTERN = re.compile(r"(?:https?://|www\.)[^\s,;)]+", re.IGNORECASE)
    _YEAR_PATTERN = re.compile(r"\b(?:19|20)\d{2}\b")
    _DURATION_PATTERN = re.compile(
        r"(?P<start>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|\d{4})\s*-\s*"
        r"(?P<end>Present|Current|Now|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|\d{4})",
        re.IGNORECASE,
    )
    _KNOWN_SKILLS = (
        "Python",
        "Java",
        "JavaScript",
        "TypeScript",
        "Node.js",
        "React",
        "Django",
        "FastAPI",
        "Flask",
        "AWS",
        "GCP",
        "Azure",
        "Docker",
        "Kubernetes",
        "SQL",
        "PostgreSQL",
        "MongoDB",
        "Redis",
        "Spark",
        "Kafka",
    )
    _SECTION_HEADERS = {
        "experience": ("experience", "work experience", "professional experience", "employment"),
        "education": ("education", "academic"),
        "projects": ("projects", "project"),
        "certifications": ("certifications", "certification", "licenses"),
        "skills": ("skills", "technical skills", "technologies"),
        "summary": ("summary", "profile", "objective"),
    }

    def __init__(
        self,
        *,
        email_normalizer: EmailNormalizer | None = None,
        phone_normalizer: PhoneNormalizer | None = None,
        skill_normalizer: SkillNormalizer | None = None,
        location_normalizer: LocationNormalizer | None = None,
        date_normalizer: DateNormalizer | None = None,
        link_classifier: LinkClassifier | None = None,
    ) -> None:
        """Initialize parser dependencies."""
        self._email_normalizer = email_normalizer or EmailNormalizer()
        self._phone_normalizer = phone_normalizer or PhoneNormalizer()
        self._skill_normalizer = skill_normalizer or SkillNormalizer()
        self._location_normalizer = location_normalizer or LocationNormalizer()
        self._date_normalizer = date_normalizer or DateNormalizer()
        self._link_classifier = link_classifier or LinkClassifier()

    def parse(self, source_path: str | Path) -> list[CandidateRecord]:
        """Parse one resume file and return a single candidate record."""
        resume_path = Path(source_path)
        logger.info("Parsing resume", extra={"source_path": str(resume_path)})

        if not resume_path.exists():
            raise ParserFileNotFoundError(f"Resume file not found: {resume_path}")

        text = self.extract_text(resume_path)
        if not text.strip():
            raise ParserSchemaError(f"Resume did not contain extractable text: {resume_path}")

        record = self._text_to_record(text, resume_path)
        logger.info("Parsed resume", extra={"source_path": str(resume_path), "candidate": record.full_name})
        return [record]

    def extract_text(self, resume_path: Path) -> str:
        """Extract text from PDF or DOCX without OCR."""
        suffix = resume_path.suffix.casefold()
        if suffix == ".pdf":
            return self._extract_pdf_text(resume_path)
        if suffix == ".docx":
            return self._extract_docx_text(resume_path)
        raise ParserSchemaError(f"Unsupported resume format: {resume_path.suffix}")

    def _extract_pdf_text(self, resume_path: Path) -> str:
        """Extract PDF text using PyMuPDF."""
        try:
            import fitz

            with fitz.open(resume_path) as document:
                return "\n".join(page.get_text() for page in document)
        except ParserSchemaError:
            raise
        except Exception as exc:
            raise ParserReadError(f"Failed to read PDF resume '{resume_path}': {exc}") from exc

    def _extract_docx_text(self, resume_path: Path) -> str:
        """Extract DOCX text using python-docx."""
        try:
            from docx import Document

            document = Document(resume_path)
            text_parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            return "\n".join(text_parts)
        except Exception as exc:
            raise ParserReadError(f"Failed to read DOCX resume '{resume_path}': {exc}") from exc

    def _text_to_record(self, text: str, resume_path: Path) -> CandidateRecord:
        """Convert extracted resume text into a candidate record."""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        sections = self._sections(lines)
        emails = self._extract_emails(text)
        phones = self._extract_phones(text)
        links = self._extract_links(text)
        skills = self._extract_skills(text, sections)
        experience = self._extract_experience(sections)
        education = self._extract_education(sections)
        projects = self._extract_projects(sections)
        certifications = self._extract_certifications(sections)
        location = self._extract_location(lines)
        summary = self._extract_summary(sections, text)

        return CandidateRecord(
            source={
                "source_type": "resume",
                "source_name": "resume",
                "source_record_id": resume_path.name,
                "source_uri": str(resume_path),
            },
            external_id=emails[0] if emails else resume_path.stem,
            full_name=self._extract_name(lines),
            emails=emails,
            phones=phones,
            location=location,
            links=links,
            headline=summary,
            skills=skills,
            experience=experience,
            education=education,
            projects=projects,
            certifications=certifications,
            resume_summary=summary,
            resume_file=resume_path.name,
            raw_values=self._raw_values(
                emails=emails,
                phones=phones,
                location=location,
                links=links,
                skills=skills,
                experience=experience,
                education=education,
                projects=projects,
                certifications=certifications,
                resume_summary=summary,
            ),
            raw_payload={"filename": resume_path.name, "text": text},
        )

    def _extract_emails(self, text: str) -> list[str]:
        """Extract and normalize email addresses."""
        return self._dedupe(
            [email for value in self._EMAIL_PATTERN.findall(text) if (email := self._email_normalizer.normalize(value))]
        )

    def _extract_phones(self, text: str) -> list[str]:
        """Extract and normalize phone numbers."""
        return self._dedupe(
            [phone for value in self._PHONE_PATTERN.findall(text) if (phone := self._phone_normalizer.normalize(value))]
        )

    def _extract_links(self, text: str) -> list[str]:
        """Extract and normalize candidate profile links."""
        links: list[str] = []
        for value in self._URL_PATTERN.findall(text):
            classified = self._link_classifier.classify(value.rstrip("."))
            if classified is not None:
                links.append(classified.url)
        return self._dedupe(links)

    def _extract_skills(self, text: str, sections: dict[str, list[str]]) -> list[str]:
        """Extract known skills from the full resume and skills section."""
        source = "\n".join(sections.get("skills", [])) or text
        found = []
        for skill in self._KNOWN_SKILLS:
            if re.search(rf"(?<![A-Za-z0-9]){re.escape(skill)}(?![A-Za-z0-9])", source, re.IGNORECASE):
                found.append(skill)
        return self._skill_normalizer.normalize_many(found)

    def _extract_experience(self, sections: dict[str, list[str]]) -> list[dict[str, Any]]:
        """Extract raw experience entries with company, title, and duration hints."""
        entries: list[dict[str, Any]] = []
        for line in sections.get("experience", []):
            if len(line) < 4:
                continue
            entry: dict[str, Any] = {"raw": line}
            duration = self._DURATION_PATTERN.search(line)
            if duration:
                entry["duration"] = duration.group(0)
                entry["start"] = self._date_normalizer.normalize(duration.group("start"))
                end_value = duration.group("end")
                entry["end"] = None if end_value.casefold() in {"present", "current", "now"} else self._date_normalizer.normalize(end_value)
            if " at " in line.casefold():
                title, company = re.split(r"\s+at\s+", line, maxsplit=1, flags=re.IGNORECASE)
                entry["title"] = title.strip(" -")
                entry["company"] = company.split(" - ")[0].strip()
            entries.append(entry)
        return entries[:10]

    def _extract_education(self, sections: dict[str, list[str]]) -> list[dict[str, Any]]:
        """Extract education entries with degree, college, and graduation year."""
        entries: list[dict[str, Any]] = []
        for line in sections.get("education", []):
            if len(line) < 3:
                continue
            entry: dict[str, Any] = {"raw": line}
            year = self._YEAR_PATTERN.search(line)
            if year:
                entry["graduation_year"] = year.group(0)
            degree = self._degree_from_line(line)
            if degree:
                entry["degree"] = degree
            institution = re.split(r",| - ", line)[0].strip()
            if institution:
                entry["institution"] = institution
            entries.append(entry)
        return entries[:5]

    def _extract_projects(self, sections: dict[str, list[str]]) -> list[dict[str, Any]]:
        """Extract project names from the projects section."""
        projects: list[dict[str, Any]] = []
        for line in sections.get("projects", []):
            if len(line) < 3:
                continue
            name = re.split(r":| - ", line, maxsplit=1)[0].strip("-* ")
            if name:
                projects.append({"name": name, "description": line})
        return projects[:10]

    def _extract_certifications(self, sections: dict[str, list[str]]) -> list[str]:
        """Extract certification names."""
        return self._dedupe([line.strip("-* ") for line in sections.get("certifications", []) if line.strip("-* ")])

    def _extract_location(self, lines: list[str]) -> str | None:
        """Extract a location-like value from early resume lines."""
        for line in lines[:8]:
            if "@" in line or "http" in line.casefold():
                continue
            if "," in line:
                location = self._location_normalizer.normalize(line)
                if location is not None:
                    return location.raw
        return None

    def _extract_summary(self, sections: dict[str, list[str]], text: str) -> str | None:
        """Extract a short resume summary."""
        summary_lines = sections.get("summary", [])
        if not summary_lines:
            return None
        summary = " ".join(summary_lines)
        return summary[:500] if summary else None

    def _extract_name(self, lines: list[str]) -> str | None:
        """Infer candidate name from the first clean resume line."""
        for line in lines[:6]:
            lowered = line.casefold()
            if "@" in line or "http" in lowered or any(char.isdigit() for char in line):
                continue
            if len(line.split()) <= 5:
                return line
        return None

    def _sections(self, lines: list[str]) -> dict[str, list[str]]:
        """Group resume lines by common section headings."""
        sections: dict[str, list[str]] = {key: [] for key in self._SECTION_HEADERS}
        current: str | None = None
        for line in lines:
            header = self._section_for_line(line)
            if header is not None:
                current = header
                continue
            if current is not None:
                sections[current].append(line)
        return sections

    def _section_for_line(self, line: str) -> str | None:
        """Return section name when a line is a known section heading."""
        normalized = line.strip().strip(":").casefold()
        for section, aliases in self._SECTION_HEADERS.items():
            if normalized in aliases:
                return section
        return None

    def _degree_from_line(self, line: str) -> str | None:
        """Extract a degree-like phrase."""
        degree_match = re.search(r"\b(B\.?S\.?|M\.?S\.?|BTech|MTech|Bachelor|Master|MBA|PhD)\b[^,;-]*", line, re.IGNORECASE)
        return degree_match.group(0).strip() if degree_match else None

    def _raw_values(self, **values: Any) -> list[dict[str, Any]]:
        """Build raw values for provenance."""
        raw_values: list[dict[str, Any]] = []
        for field_name, value in values.items():
            if value in (None, "", [], {}):
                continue
            raw_values.append({"field_name": field_name, "source_key": field_name, "value": value})
        return raw_values

    def _dedupe(self, values: list[str]) -> list[str]:
        """Deduplicate strings while preserving order."""
        deduped: list[str] = []
        seen: set[str] = set()
        for value in values:
            key = value.casefold()
            if key not in seen:
                deduped.append(value)
                seen.add(key)
        return deduped
